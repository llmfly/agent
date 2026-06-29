#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成DOCX格式的论文解析报告"""

import sys
import subprocess
import os

# 确保python-docx已安装
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document

from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 设置默认字体 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ===== 辅助函数 =====
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, size=None, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

# =============================================
# 封面
# =============================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('论文解析报告')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run(
    'Expansion-Squeeze-Excitation Fusion Network\n'
    'for Elderly Activity Recognition'
)
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run(
    'IEEE Transactions on Circuits and Systems for Video Technology\n'
    'Vol. 32, No. 8, August 2022'
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('解析日期：2026年6月26日')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# =============================================
# 一、基本信息
# =============================================
add_heading_styled('一、基本信息', 1)

add_table(
    ['项目', '内容'],
    [
        ['论文标题', 'Expansion-Squeeze-Excitation Fusion Network for Elderly Activity Recognition'],
        ['中文译名', '面向老年人活动识别的扩展-压缩-激励融合网络'],
        ['作者', 'Xiangbo Shu, Jiawen Yang, Rui Yan, Yan Song'],
        ['单位', '南京理工大学 计算机科学与工程学院'],
        ['发表期刊', 'IEEE TCSVT, Vol. 32, No. 8, Aug. 2022'],
        ['页数', '12页（页码 5281-5292）'],
        ['投稿/录用日期', '2021-11-16投稿 / 2022-01-09录用'],
    ]
)
doc.add_paragraph()

# =============================================
# 二、研究背景与动机
# =============================================
add_heading_styled('二、研究背景与动机', 1)

add_heading_styled('2.1 问题定义', 2)
add_para(
    '本文聚焦于老年人活动识别（Elderly Activity Recognition）任务。'
    '随着人口老龄化加剧，空巢老人比例不断上升，利用计算机视觉技术对老年人日常活动进行智能监测具有重要意义。'
)

add_heading_styled('2.2 核心挑战', 2)
add_para('与常规动作识别相比，老年人活动识别面临以下独特挑战：')

challenges = [
    ('个体动作与人-物交互并存', '老年人活动中既包含纯身体动作，也包含与物体交互的行为（如吹头发、梳头）。'),
    ('局部交互微小', '许多人-物交互动作是局部性的（如手部细微动作），容易被忽略。'),
    ('动作幅度不明显', '老年人动作幅度较小，特征不显著。'),
    ('相似动作混淆', '某些动作在时空轨迹上高度相似，仅靠微小局部差异才能区分。'),
]
for title, desc in challenges:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

add_para('')
p = doc.add_paragraph(style='List Bullet')
p.add_run('典型案例：')
p.add_run('"吹头发" vs "梳头"（交互物体不同）、"打电话" vs "玩手机"（局部手部动作差异）')

# =============================================
# 三、核心方法
# =============================================
add_heading_styled('三、核心方法：ESE-FN', 1)

add_heading_styled('3.1 总体框架', 2)
add_para('ESE-FN（Expansion-Squeeze-Excitation Fusion Network）的整体框架包含四个核心模块：')

modules = [
    '特征提取模块（Feature Extractor Module）：RGB → ResNeXt101；骨架 → Shift-GCN',
    '模态融合模块（M-Net）：通过模态级ESE注意力实现模态间粗粒度融合',
    '通道融合模块（C-Net）：通过通道级ESE注意力实现通道间细粒度融合',
    '多模态损失（Multi-modal Loss, ML）：保持单模态与融合特征的一致性',
]
for m in modules:
    doc.add_paragraph(m, style='List Number')

add_para('')
add_para('整体流程示意：', bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'RGB视频 → ResNeXt101 → RGB特征 fr\n'
    '骨架序列 → Shift-GCN → 骨架特征 fs\n'
    '        ↓\n'
    '特征拼接 → M-Net（模态级ESE融合）\n'
    '         → C-Net（通道级ESE融合）→ 融合特征 frs\n'
    '        ↓\n'
    '多模态损失 L = α×L_rs + β×(min(L_r,L_s)-L_rs)'
)
run.font.size = Pt(10)

add_heading_styled('3.2 SENet回顾与改进', 2)
add_para(
    'SENet通过Squeeze（全局平均池化）和Excitation（FC+Sigmoid）两步学习通道级注意力。'
    '本文的创新在于引入Expansion（扩展）步骤，形成ESE结构，实现从局部和全局两个视角特征交互。'
)

add_heading_styled('3.3 M-Net（模态融合网络）', 2)
add_para('M-Net通过M-ESEA实现模态间粗粒度融合的三个步骤：')

p = doc.add_paragraph(style='List Number')
run = p.add_run('模态级扩展：')
run.bold = True
p.add_run('多个不同卷积核的卷积层堆叠，从局部视角扩展模态信息')

p = doc.add_paragraph(style='List Number')
run = p.add_run('模态级压缩：')
run.bold = True
p.add_run('全局平均池化，从全局视角获得模态级表示')

p = doc.add_paragraph(style='List Number')
run = p.add_run('模态级激励：')
run.bold = True
p.add_run('FC层+Sigmoid学习M-ESEA注意力权重，更新特征')

add_heading_styled('3.4 C-Net（通道融合网络）', 2)
add_para('C-Net通过C-ESEA实现通道间细粒度融合的三个步骤：')

p = doc.add_paragraph(style='List Number')
run = p.add_run('通道级扩展：')
run.bold = True
p.add_run('单层卷积扩展通道信息（轻量化设计）')

p = doc.add_paragraph(style='List Number')
run = p.add_run('通道级压缩：')
run.bold = True
p.add_run('全局平均池化获得通道级表示')

p = doc.add_paragraph(style='List Number')
run = p.add_run('通道级激励：')
run.bold = True
p.add_run('FC层+Sigmoid学习C-ESEA注意力权重')

add_para('')
add_para('设计思路：M-Net（粗融合，多层卷积）→ C-Net（细融合，单层卷积），逐步递进。')

add_heading_styled('3.5 多模态损失（ML）', 2)
add_para('公式：L = α × L_rs + β × (min(L_r, L_s) − L_rs)')
add_para('创新点：引入惩罚项，促使融合特征的预测损失不低于最优单模态。超参数最优值 α=0.7, β=0.3。')

# =============================================
# 四、实验设置
# =============================================
add_heading_styled('四、实验设置', 1)

add_heading_styled('4.1 数据集', 2)
add_para('主要：ETRI-Activity3D（最大老年人活动识别数据集，112,620样本，55类，100人）')
add_para('扩展：NTU RGB+D（56,880骨架序列，60类，40人，CS和CV两种协议）')

add_heading_styled('4.2 实现细节', 2)
add_table(
    ['参数', '设置'],
    [
        ['RGB主干', 'ResNeXt101（默认）/ ResNeXt18（消融）'],
        ['骨架主干', 'Shift-GCN'],
        ['视频分片', 'T=64 clips'],
        ['Batch Size', '32'],
        ['Epochs', '30（ESE-FN训练）'],
        ['优化器', 'SGD（动量0.9，学习率0.1）'],
        ['框架/硬件', 'PyTorch / Titan RTX GPU'],
    ]
)

# =============================================
# 五、实验结果
# =============================================
add_heading_styled('五、实验结果', 1)

add_heading_styled('5.1 消融实验', 2)
add_para('组件消融：M-Net、C-Net、ML三个组件均正向贡献，完整ESE-FN最优。')
add_para('与SENet对比：ESE-FN比SENet高出0.6%，验证Expansion步骤有效性。')
add_para('双重融合策略（模态+通道）优于单一融合策略。')

add_heading_styled('5.2 老年人活动识别（ETRI-Activity3D）', 2)
add_table(
    ['方法', '准确率'],
    [
        ['FSA-CNN (前SOTA)', '93.7%'],
        ['ESE-FN (Ours)', '95.9% ✦'],
    ]
)
add_para('✦ SOTA，比前SOTA高出2.2%（相对提升2.3%）')

add_heading_styled('5.3 常规动作识别（NTU RGB+D）', 2)
add_table(
    ['方法', 'CS', 'CV'],
    [
        ['Evolution Pose Map', '91.7%', '-'],
        ['Shift-GCN', '-', '96.5%'],
        ['ESE-FN (Ours)', '92.4%', '96.7% ✦'],
    ]
)
add_para('✦ CS提升0.7%，CV提升0.2%，验证泛化能力')

# =============================================
# 六、贡献总结
# =============================================
add_heading_styled('六、论文贡献总结', 1)

contributions = [
    ('提出ESE-FN框架', '首次将Expansion引入SENet家族，形成ESE注意力结构用于多模态特征融合'),
    ('设计M-Net和C-Net', '分别从模态级和通道级捕获特征依赖，实现从粗到细的递进式融合'),
    ('提出多模态损失ML', '通过惩罚项保持单模态与融合特征之间的一致性'),
    ('取得SOTA性能', 'ETRI-Activity3D上95.9%，NTU RGB+D上验证泛化能力'),
]
for title, desc in contributions:
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

# =============================================
# 七、论文结构
# =============================================
add_heading_styled('七、论文结构', 1)
add_table(
    ['章节', '内容'],
    [
        ['I. Introduction', '研究背景、问题定义、挑战、贡献总结'],
        ['II. Related Work', 'RGB/骨架/多模态融合方法综述'],
        ['III. Methodology', 'SENet回顾、ESE-FN框架、M-Net/C-Net/ML详述'],
        ['IV. Experiments', '数据集、实验细节、消融实验、对比实验'],
        ['V. Conclusion', '方法总结、贡献重申、讨论与展望'],
    ]
)

# =============================================
# 八、讨论
# =============================================
add_heading_styled('八、讨论与展望', 1)
add_para(
    '本文在SENet基础上首次引入Expansion步骤，实现上采样和下采样的特征交互，'
    '是SENet家族中的新尝试。ESE-FN不仅限于老年人活动识别，'
    '还可推广到一般特征学习和多模态特征融合任务，为后续研究提供了新的思路。'
)

# ===== 尾部 =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告完 —')
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('生成日期：2026-06-26 | 源文件：IEEE TCSVT 2022')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ===== 保存 =====
output_path = '/data/intelli/engine/.deer-flow/users/25ff167c-03fd-4d28-a5e7-52bea757d603/threads/65882f9e-4f7e-40cf-bf99-678786dbb633/user-data/outputs/解析报告_ESE-FN论文.docx'
doc.save(output_path)
print(f'SUCCESS: 报告已保存至 {output_path}')
