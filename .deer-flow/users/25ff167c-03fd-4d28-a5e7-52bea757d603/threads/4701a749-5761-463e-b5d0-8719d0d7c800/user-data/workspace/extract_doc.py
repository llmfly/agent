#!/usr/bin/env python3
"""Extract text from docx file."""
from docx import Document
import sys

filepath = sys.argv[1]
doc = Document(filepath)

print("=== 段落内容 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[段落 {i}] {para.text}")

print("\n=== 表格内容 ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- 表格 {ti+1} ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"行 {ri}: {' | '.join(cells)}")

print("\n=== 文档摘要 ===")
print(f"段落总数: {len(doc.paragraphs)}")
print(f"非空段落数: {sum(1 for p in doc.paragraphs if p.text.strip())}")
print(f"表格数: {len(doc.tables)}")

# Also try to read headers/footers
for section in doc.sections:
    header = section.header
    footer = section.footer
    if header.paragraphs:
        for p in header.paragraphs:
            if p.text.strip():
                print(f"页眉: {p.text}")
    if footer.paragraphs:
        for p in footer.paragraphs:
            if p.text.strip():
                print(f"页脚: {p.text}")
