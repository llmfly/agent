"""PDF/DOCX parsing tools for Lead Agent direct invocation.

These are thin wrappers around the Worker implementations.
Lead Agent calls these Tools; the Execution Layer uses the Workers directly.
"""

from __future__ import annotations

import logging

from langchain_core.tools import tool

from deerflow.tools.types import Runtime

logger = logging.getLogger(__name__)


def _resolve_file_path(file_path: str, runtime: Runtime) -> str | None:
    """Resolve a sandbox virtual path to the host filesystem path.

    The Tool runs in the backend process (not in the sandbox), so sandbox
    virtual paths like ``/mnt/user-data/uploads/file.pdf`` must be resolved
    to the actual host path via ``DeerFlowPaths.resolve_virtual_path()``.
    """
    import os
    from pathlib import Path

    # 1. Direct absolute path — already on the host filesystem
    if os.path.isabs(file_path) and os.path.exists(file_path):
        logger.debug("_resolve_file_path: direct hit %s", file_path)
        return file_path

    # 2. Extract thread_id from runtime context
    thread_id = None
    user_id = None
    try:
        if runtime.context:
            thread_id = runtime.context.get("thread_id")
            user_id = runtime.context.get("user_id")
        if not thread_id:
            cfg = getattr(runtime, "config", None) or {}
            thread_id = cfg.get("configurable", {}).get("thread_id")
        if not thread_id:
            try:
                from langgraph.config import get_config
                thread_id = get_config().get("configurable", {}).get("thread_id")
            except RuntimeError:
                pass
    except Exception:
        pass

    basename = os.path.basename(file_path)
    paths = None

    # 3. Try resolving via resolve_virtual_path (current thread)
    try:
        if thread_id:
            from deerflow.config.paths import get_paths
            paths = get_paths()
            resolved = paths.resolve_virtual_path(thread_id, file_path)
            if resolved.exists():
                logger.debug("_resolve_file_path: virtual_path hit for thread %s", thread_id)
                return str(resolved)
    except Exception:
        pass

    # 4. Fallback: check current thread's uploads dir directly
    try:
        if thread_id and paths:
            uploads_dir = paths.sandbox_uploads_dir(thread_id)
            candidate = uploads_dir / basename
            if candidate.exists():
                logger.debug("_resolve_file_path: uploads_dir hit for thread %s", thread_id)
                return str(candidate)
    except Exception:
        pass

    # 5. Cross-thread fallback: search ALL user threads for this filename.
    #    This handles the case where the file was uploaded in a different
    #    thread than the current conversation.
    try:
        if not paths:
            from deerflow.config.paths import get_paths
            paths = get_paths()

        # Resolve the user's threads base directory
        if user_id:
            threads_root = Path(paths.user_dir(user_id)) / "threads"
        elif thread_id:
            # Derive user from the thread dir parent structure
            thread_dir = Path(paths.sandbox_uploads_dir(thread_id)).parent.parent.parent
            threads_root = thread_dir.parent  # {user_dir}/threads/
        else:
            threads_root = None

        if threads_root and threads_root.exists():
            for thread_dir in threads_root.iterdir():
                candidate = thread_dir / "user-data" / "uploads" / basename
                if candidate.exists():
                    logger.debug(
                        "_resolve_file_path: cross-thread hit %s in thread %s",
                        basename, thread_dir.name,
                    )
                    return str(candidate)
    except Exception:
        pass

    logger.warning("_resolve_file_path: file NOT found: %s (thread=%s)", file_path, thread_id)
    return None


@tool(parse_docstring=True)
def parse_pdf(
    file_path: str,
    runtime: Runtime,
    max_pages: int = 0,
) -> str:
    """解析 PDF 文件，提取文本内容和页数信息。

    使用 PyMuPDF 库解析 PDF 文件，返回每页文本。
    当用户上传 PDF 文件并要求解析内容时，必须调用此工具。

    Args:
        file_path: PDF 文件的路径。
        max_pages: 最大解析页数。0 表示全部。默认 0。
    """
    import fitz
    import os

    resolved = _resolve_file_path(file_path, runtime)
    if not resolved:
        return f"错误：找不到文件 '{file_path}'。"

    try:
        doc = fitz.open(resolved)
        total_pages = len(doc)
        pages_to_read = total_pages if max_pages <= 0 else min(max_pages, total_pages)

        parts = [f"📄 文档信息: 共 {total_pages} 页", f"文件: {os.path.basename(resolved)}", ""]
        for i in range(pages_to_read):
            text = doc[i].get_text().strip()
            if text:
                if len(text) > 8000:
                    text = text[:8000] + f"\n... (已截断, 原长度 {len(text)} 字符)"
                parts.append(f"--- 第 {i+1} 页 ---")
                parts.append(text)
                parts.append("")

        if pages_to_read < total_pages:
            parts.append(f"... 共 {total_pages} 页，仅展示前 {pages_to_read} 页")

        doc.close()
        return "\n".join(parts)
    except Exception as e:
        logger.exception("PDF 解析失败")
        return f"错误：PDF 解析失败 - {e}"


@tool(parse_docstring=True)
def parse_docx(
    file_path: str,
    runtime: Runtime,
) -> str:
    """解析 DOCX 文件，提取文本内容和段落结构。

    使用 python-docx 库解析 Word 文档，返回文本段落、标题层级和表格。
    当用户上传 Word 文件并要求解析内容时，必须调用此工具。

    Args:
        file_path: DOCX 文件的路径。
    """
    from docx import Document
    import os

    resolved = _resolve_file_path(file_path, runtime)
    if not resolved:
        return f"错误：找不到文件 '{file_path}'。"

    try:
        doc = Document(resolved)
        parts = [f"📄 文档信息: {os.path.basename(resolved)}", ""]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading" in style or "title" in style:
                level = style.replace("heading ", "").replace("headings ", "")
                prefix = "#" * (int(level) if level.isdigit() else 1)
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

        parts.append("")
        if doc.tables:
            parts.append(f"--- 表格 ({len(doc.tables)} 个) ---")
            for ti, table in enumerate(doc.tables):
                parts.append(f"\n表格 {ti+1}:")
                for ri, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(cells))
                    if ri == 0:
                        parts.append("-" * 60)

        result = "\n".join(parts)
        if len(result) > 50000:
            result = result[:50000] + "\n... (已截断)"

        return result
    except Exception as e:
        logger.exception("DOCX 解析失败")
        return f"错误：DOCX 解析失败 - {e}"


@tool(parse_docstring=True)
def extract_pdf_metadata(
    file_path: str,
    runtime: Runtime,
) -> str:
    """提取 PDF 文件的元数据信息（标题、作者、页数等）。

    不提取正文内容，适合快速了解文档概况。

    Args:
        file_path: PDF 文件的路径。
    """
    import fitz
    import os

    resolved = _resolve_file_path(file_path, runtime)
    if not resolved:
        return f"错误：找不到文件 '{file_path}'。"

    try:
        doc = fitz.open(resolved)
        meta = doc.metadata or {}
        has_text = any(page.get_text().strip() for page in doc)
        lines = [
            f"📄 PDF 元数据: {os.path.basename(resolved)}",
            f"  页数: {len(doc)}",
            f"  标题: {meta.get('title', 'N/A')}",
            f"  作者: {meta.get('author', 'N/A')}",
            f"  主题: {meta.get('subject', 'N/A')}",
            f"  关键词: {meta.get('keywords', 'N/A')}",
            f"  创建时间: {meta.get('creationDate', 'N/A')}",
            f"  修改时间: {meta.get('modDate', 'N/A')}",
            f"  PDF 版本: {doc.pdf_version}",
            f"  包含可提取文本: {'是' if has_text else '否'}",
        ]
        doc.close()
        return "\n".join(lines)
    except Exception as e:
        return f"错误：元数据提取失败 - {e}"
