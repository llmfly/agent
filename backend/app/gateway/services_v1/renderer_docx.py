"""DOCX report renderer.

Converts a ReportSpec into a .docx file using the python-docx library.
"""

from __future__ import annotations

import re
from datetime import datetime
from typing import Any

from app.gateway.schemas.v1.reports import ContentBlock, ReportSection, ReportSpec, TableContent
from app.gateway.services_v1.renderer_base import BaseRenderer

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxRenderer(BaseRenderer):
    """Render ReportSpec to a .docx file."""

    def __init__(self) -> None:
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX rendering. "
                "Install it with: uv add python-docx"
            )

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self) -> str:
        return "docx"

    @staticmethod
    def _sanitize(text: str) -> str:
        """Strip NULL bytes and XML-incompatible control characters.

        XML 1.0 allows: tab (0x09), LF (0x0A), CR (0x0D), and
        0x20–0xD7FF / 0xE000–0xFFFD.  Everything else is stripped.
        """
        return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)

    @staticmethod
    def _clean(p: Any, text: str, **kwargs: Any) -> Any:
        """Add a sanitized run to a paragraph."""
        return p.add_run(DocxRenderer._sanitize(text), **kwargs)

    @staticmethod
    def _para(doc: Document, text: str = "", **kwargs: Any) -> Any:
        """Add a paragraph with sanitized text."""
        return doc.add_paragraph(DocxRenderer._sanitize(text), **kwargs)

    def render(self, spec: ReportSpec) -> bytes:
        """Render ReportSpec to a DOCX document."""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        self._add_title_page(doc, spec)
        self._add_toc(doc, spec)
        self._add_sections(doc, spec)

        if spec.citations:
            self._add_citations(doc, spec.citations)

        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = self._clean(p, f"由 intelli-engine AI 能力平台生成 | {datetime.now().isoformat()}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Save to bytes
        from io import BytesIO
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def _add_title_page(self, doc: Document, spec: ReportSpec) -> None:
        """Add the title page."""
        # Empty paragraphs for spacing
        for _ in range(4):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = self._clean(title, spec.title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        if spec.subtitle:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = self._clean(subtitle, spec.subtitle)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = self._clean(meta,
            f"作者: {spec.metadata.get('author', 'intelli-engine')}  |  "
            f"语言: {spec.metadata.get('language', 'zh-CN')}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

    def _add_toc(self, doc: Document, spec: ReportSpec) -> None:
        """Add a table of contents page."""
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = self._clean(heading, "目录")
        run.font.size = Pt(18)
        run.font.bold = True

        doc.add_paragraph()

        for i, section in enumerate(spec.sections, 1):
            p = doc.add_paragraph()
            run = self._clean(p, f"  {i}.  {section.heading}")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

        doc.add_page_break()

    def _add_sections(self, doc: Document, spec: ReportSpec) -> None:
        """Add all report sections."""
        for section in spec.sections:
            self._add_section_heading(doc, section.heading)
            for block in section.content:
                self._add_content_block(doc, block)
            doc.add_paragraph()  # spacing

    def _add_section_heading(self, doc: Document, heading: str) -> None:
        """Add a section heading."""
        p = doc.add_paragraph()
        run = self._clean(p, heading)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '4',
            qn('w:color'): 'CCCCCC',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_paragraph()  # spacing

    def _add_content_block(self, doc: Document, block: ContentBlock) -> None:
        """Add a content block to the document."""
        if block.type == "paragraph" and block.text:
            p = self._para(doc, block.text)
            p.paragraph_format.space_after = Pt(6)

        elif block.type == "heading":
            level = min(block.level or 2, 6)
            p = doc.add_paragraph()
            run = self._clean(p, block.text or "")
            run.font.size = Pt({2: 14, 3: 13, 4: 12}.get(level, 15))
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif block.type == "bullets" and block.items:
            for item in block.items:
                p = self._para(doc, item, style='List Bullet')
                p.paragraph_format.space_after = Pt(3)

        elif block.type == "numbered_list" and block.items:
            for i, item in enumerate(block.items, 1):
                p = self._para(doc, f"{i}. {item}")
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.5)

        elif block.type == "table" and block.table:
            self._add_table(doc, block.table)

        elif block.type == "code" and block.code:
            p = doc.add_paragraph()
            run = self._clean(p, block.code)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
            p.paragraph_format.left_indent = Inches(0.3)
            # Add shading
            shading = p._p.get_or_add_pPr().makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F5F5F5',
            })
            p._p.get_or_add_pPr().append(shading)

        elif block.type == "quote" and block.text:
            p = self._para(doc, block.text)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0] if p.runs else self._clean(p, block.text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_table(self, doc: Document, table_content: TableContent) -> None:
        """Add a table to the document."""
        num_rows = len(table_content.rows) + 1
        num_cols = len(table_content.columns)

        if num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for j, col_name in enumerate(table_content.columns):
            cell = table.rows[0].cells[j]
            cell.text = self._sanitize(col_name)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for i, row in enumerate(table_content.rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.rows[i + 1].cells[j]
                    cell.text = self._sanitize(str(cell_text))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()  # spacing

    def _add_citations(self, doc: Document, citations: list[Any]) -> None:
        """Add citations section."""
        doc.add_page_break()
        self._add_section_heading(doc, "参考资料")

        for c in citations:
            p = doc.add_paragraph()
            run = self._clean(p, f"[{c.id}] {c.label}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if c.locator:
                run2 = self._clean(p, f" — {c.locator}")
                run2.font.size = Pt(10)
                run2.font.italic = True
                run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
