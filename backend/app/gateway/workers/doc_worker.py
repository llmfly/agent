"""Document parsing Worker — PDF and DOCX extraction for the Execution Layer.

Registered as capability ``document_parse`` in CapabilityRegistry.
The Execution Planner dispatches Business Tasks involving document
analysis to this Worker. Output is always ``List[Evidence]``.
"""

from __future__ import annotations

import logging
import os
from typing import Any

from app.gateway.models.evidence import Content, Evidence, EvidenceType, SourceInfo
from app.gateway.workers.base import BaseWorker, ExecutionTask

logger = logging.getLogger(__name__)


def _resolve_file_path(file_path: str) -> str | None:
    """Resolve a file path against known workspace locations."""
    if os.path.isabs(file_path) and os.path.exists(file_path):
        return file_path

    sandbox_root = os.environ.get("SANDBOX_ROOT") or "/mnt/user-data"
    for base in [sandbox_root]:
        candidate = os.path.join(base, file_path.lstrip("/"))
        if os.path.exists(candidate):
            return candidate

    return None


class PdfWorker(BaseWorker):
    """Extract text from PDF files using PyMuPDF."""

    name: str = "pdf"
    capability: str = "document_parse"

    async def execute(self, task: ExecutionTask, context: dict[str, Any]) -> list[Evidence]:
        file_path = task.params.get("file_path", "")
        max_pages: int = task.params.get("max_pages", 0)

        # Prefer the pre-resolved path from pipeline context
        resolved = context.get("resolved_doc_path", "")
        if not resolved or not os.path.exists(resolved):
            resolved = _resolve_file_path(file_path)
        if not resolved:
            return [
                Evidence(
                    type="pdf_chunk",
                    source=SourceInfo(datasource_type="pdf", file=file_path),
                    content=Content(text=f"错误：找不到文件 '{file_path}'"),
                    score=0.0,
                )
            ]

        import fitz

        try:
            doc = fitz.open(resolved)
            total_pages = len(doc)
            pages_to_read = total_pages if max_pages <= 0 else min(max_pages, total_pages)
            results: list[Evidence] = []

            for i in range(pages_to_read):
                page = doc[i]
                text = page.get_text().strip()
                if not text:
                    continue

                ev = Evidence(
                    id=f"pdf_{os.path.basename(resolved)}_{i+1}",
                    type="pdf_chunk",
                    source=SourceInfo(
                        datasource_type="pdf",
                        file=resolved,
                        document_id=os.path.basename(resolved),
                    ),
                    content=Content(text=text[:8000]),
                    metadata={
                        "page": i + 1,
                        "total_pages": total_pages,
                        "filename": os.path.basename(resolved),
                    },
                    score=1.0,
                )
                results.append(ev)

            doc.close()
            return results

        except Exception as e:
            logger.exception("PDFWorker 解析失败")
            return [
                Evidence(
                    type="pdf_chunk",
                    source=SourceInfo(datasource_type="pdf", file=file_path),
                    content=Content(text=f"PDF 解析失败: {e}"),
                    score=0.0,
                )
            ]

    async def validate(self, task: ExecutionTask) -> list[str]:
        warnings: list[str] = []
        if "file_path" not in task.params:
            warnings.append("缺少 file_path 参数")
        return warnings


class DocxWorker(BaseWorker):
    """Extract text from DOCX files using python-docx."""

    name: str = "docx"
    capability: str = "document_parse"

    async def execute(self, task: ExecutionTask, context: dict[str, Any]) -> list[Evidence]:
        file_path = task.params.get("file_path", "")

        resolved = context.get("resolved_doc_path", "")
        if not resolved or not os.path.exists(resolved):
            resolved = _resolve_file_path(file_path)
        if not resolved:
            return [
                Evidence(
                    type="docx_chunk",
                    source=SourceInfo(datasource_type="docx", file=file_path),
                    content=Content(text=f"错误：找不到文件 '{file_path}'"),
                    score=0.0,
                )
            ]

        from docx import Document

        try:
            doc = Document(resolved)
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style = para.style.name.lower() if para.style else ""
                    if "heading" in style or "title" in style:
                        level = style.replace("heading ", "").replace("headings ", "")
                        prefix = "#" * (int(level) if level.isdigit() else 1)
                        paragraphs.append(f"{prefix} {text}")
                    else:
                        paragraphs.append(text)

            full_text = "\n".join(paragraphs)
            if len(full_text) > 50000:
                full_text = full_text[:50000] + f"\n... (已截断, 原长度 {len(full_text)} 字符)"

            ev = Evidence(
                type="docx_chunk",
                source=SourceInfo(
                    datasource_type="docx",
                    file=resolved,
                    document_id=os.path.basename(resolved),
                ),
                content=Content(text=full_text),
                metadata={
                    "filename": os.path.basename(resolved),
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                },
                score=1.0,
            )

            # Also extract tables as separate evidence
            results = [ev]
            for ti, table in enumerate(doc.tables):
                rows_text = []
                for row in table.rows:
                    rows_text.append(" | ".join(cell.text.strip() for cell in row.cells))
                table_text = "\n".join(rows_text)
                if table_text.strip():
                    results.append(
                        Evidence(
                            type="docx_chunk",
                            source=SourceInfo(
                                datasource_type="docx",
                                file=resolved,
                                document_id=os.path.basename(resolved),
                            ),
                            content=Content(text=f"[表格 {ti+1}]\n{table_text}"),
                            metadata={"table_index": ti, "rows": len(table.rows), "cols": len(table.columns)},
                            score=0.9,
                        )
                    )

            return results

        except Exception as e:
            logger.exception("DocxWorker 解析失败")
            return [
                Evidence(
                    type="docx_chunk",
                    source=SourceInfo(datasource_type="docx", file=file_path),
                    content=Content(text=f"DOCX 解析失败: {e}"),
                    score=0.0,
                )
            ]

    async def validate(self, task: ExecutionTask) -> list[str]:
        warnings: list[str] = []
        if "file_path" not in task.params:
            warnings.append("缺少 file_path 参数")
        return warnings


class PdfMetadataWorker(BaseWorker):
    """Extract PDF metadata (title, author, pages, etc.) without full text."""

    name: str = "pdf_metadata"
    capability: str = "document_parse"

    async def execute(self, task: ExecutionTask, context: dict[str, Any]) -> list[Evidence]:
        file_path = task.params.get("file_path", "")

        resolved = context.get("resolved_doc_path", "")
        if not resolved or not os.path.exists(resolved):
            resolved = _resolve_file_path(file_path)
        if not resolved:
            return [
                Evidence(
                    type="pdf_chunk",
                    source=SourceInfo(datasource_type="pdf", file=file_path),
                    content=Content(text=f"错误：找不到文件 '{file_path}'"),
                    score=0.0,
                )
            ]

        import fitz

        try:
            doc = fitz.open(resolved)
            meta = doc.metadata or {}
            has_text = any(page.get_text().strip() for page in doc)

            lines = [
                f"📄 PDF 元数据: {os.path.basename(resolved)}",
                f"  页数: {len(doc)}",
                f"  标题: {meta.get('title', 'N/A')}",
                f"  作者: {meta.get('author', 'N/A')}",
                f"  主题: {meta.get('subject', 'N/A')}",
                f"  关键词: {meta.get('keywords', 'N/A')}",
                f"  创建时间: {meta.get('creationDate', 'N/A')}",
                f"  修改时间: {meta.get('modDate', 'N/A')}",
                f"  PDF 版本: {doc.pdf_version}",
                f"  包含可提取文本: {'是' if has_text else '否'}",
            ]
            doc.close()

            return [
                Evidence(
                    type="pdf_chunk",
                    source=SourceInfo(datasource_type="pdf", file=resolved),
                    content=Content(text="\n".join(lines)),
                    metadata={
                        "filename": os.path.basename(resolved),
                        "pages": len(doc),
                        "has_text": has_text,
                    },
                    score=1.0,
                )
            ]

        except Exception as e:
            return [
                Evidence(
                    type="pdf_chunk",
                    source=SourceInfo(datasource_type="pdf", file=file_path),
                    content=Content(text=f"元数据提取失败: {e}"),
                    score=0.0,
                )
            ]
